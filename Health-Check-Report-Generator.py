import json
from docx import Document
from docx.shared import RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googletrans import Translator

# Crear una instancia del objeto Translator
translator = Translator()

# Función para cargar archivos JSON
def cargar_json(ruta):
    with open(ruta) as archivo:
        return json.load(archivo)

# Función para seleccionar una opción del menú
def seleccionar_opcion(datos):
    print("📝Selecciona una opción:")
    for idx, data in enumerate(datos):
        print(f"{idx}: {data['runOn']}")
    opcion = int(input("Opción: "))
    if opcion < 0 or opcion >= len(datos):
        raise ValueError(" ❌Opción inválida")
    return opcion

# Función para crear la portada del documento
def crear_portada(document, cliente):
    print('🖼️  Creando Documento')
    document.add_heading('Informe de estado de la plataforma', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading('Plataforma: Trellix EPO', level=4)
    document.add_heading(f'Cliente: {cliente}', level=4)
    document.add_heading('Fecha: ', level=4)
    document.add_picture('portada.png', width=Inches(8))
    document.add_page_break()

# Función para generar el análisis de cada ítem del Health Check
def generar_analisis(document, run_on, translator):
    successful, failed, warning = 0, 0, 0
    print('🤖 Traduciciendo')
    for item in run_on["checkResultList"][:15]:  # Limitar a los primeros 15 elementos
        item_table = document.add_table(rows=1, cols=2)
        item_table.style = 'Light List Accent 1'

        # Título y estado
        item_title_cell = item_table.cell(0, 0)
        item_title_cell.text = item['checkName']
        item_title_cell.paragraphs[0].runs[0].bold = True

        item_status_cell = item_table.cell(0, 1)
        item_status_cell.text = f"Estado: {item['status']}"
        item_status_cell.paragraphs[0].runs[0].bold = Trues

        # Contabilizar los estados
        status = item["status"]
        if status == "Successful":
            successful += 1
        elif status == "Failed":
            failed += 1
        elif status == "Warning":
            warning += 1

        # Descripción traducida
        
        item_description = item['checkDescription'].replace("<br/>", "")
        item_description_es = translator.translate(item_description, dest='es').text
        document.add_paragraph(item_description_es)

        # Procesar los dataInfoList si existen
        if item.get("dataInfoList"):
            for data in item["dataInfoList"]:
                indicator_table = document.add_table(rows=1, cols=2)
                name_es = translator.translate(data["nameInfoItem"]["renderingData"], dest='es').text
                value_es = translator.translate(data["valueInfoItem"]["renderingData"], dest='es').text
                indicator_table.cell(0, 0).text = name_es
                indicator_table.cell(0, 1).text = value_es
    
    return successful, failed, warning

# Función para generar el resumen de estado
def generar_resumen_estado(document, successful, failed, warning):
    document.add_heading('Resumen de hallazgos ', level=3)
    status_table = document.add_table(rows=4, cols=2, style='Medium Shading 2 Accent 1')

    estados = ['Estado', 'Fallido', 'Advertencia', 'Exitosa']
    valores = [None, failed, warning, successful]
    colores = [None, RGBColor(255, 0, 0), RGBColor(226, 135, 67), RGBColor(0, 204, 0)]

    for i in range(4):
        status_table.cell(i, 0).text = estados[i]
        if i > 0:
            paragraph = status_table.cell(i, 1).add_paragraph()
            run = paragraph.add_run(str(valores[i]))
            run.font.color.rgb = colores[i]

# Función para generar la información del servidor
def generar_info_servidor(document, info, translator):
    document.add_heading('Información del servidor', level=2)
    categorias = ["Información del servidor", "Información EPO Trellix", "Información de servidor SQL", "Información de la base de datos"]
    
    for idx, categoria in enumerate(categorias):
        document.add_heading(categoria, level=3)
        for item in info[idx]['serverInfoItems']:
            table = document.add_table(rows=1, cols=2, style='TableGrid')
            name_es = translator.translate(item['nameInfoItem']['renderingData'], dest='es').text
            value = item['valueInfoItem']['renderingData']
            table.cell(0, 0).text = name_es
            table.cell(0, 1).text = value

# Función principal
def generar_informe(cliente, ruta_resultados, ruta_serverinfo):
    document = Document()
    document.settings.odd_and_even_pages_header_footer = True

    # Cargar datos
    datos = cargar_json(ruta_resultados)
    opcion = seleccionar_opcion(datos)
    run_on = datos[opcion]
    info = cargar_json(ruta_serverinfo)

    # Crear informe
    crear_portada(document, cliente)
    successful, failed, warning = generar_analisis(document, run_on, translator)
    generar_resumen_estado(document, successful, failed, warning)
    generar_info_servidor(document, info, translator)

    # Guardar informe
    document.save(f'salida-{cliente}.docx')
    print(f"Informe generado para {cliente}")

# Ejecución del script
if __name__ == '__main__':
    print('🚀 Trellix EPO Health Check Report Generator')
    cliente = input('👨‍💻 ¿A qué cliente pertenece? ')
    ruta_resultados = 'healthcheckresults.txt'
    ruta_serverinfo = 'serverinfo.txt'
    
    generar_informe(cliente, ruta_resultados, ruta_serverinfo)
